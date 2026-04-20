from __future__ import annotations

import mimetypes
import re
import zipfile
from collections.abc import Iterable
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from .base import DocumentSection, KnowledgeAsset, ParsedDocument, ParsedDocumentAdapter


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _tokenize_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines()]


class TextAdapter(ParsedDocumentAdapter):
    kind = "text"
    supported_suffixes = (".txt", ".rst", ".log")

    async def parse(self, asset: KnowledgeAsset) -> ParsedDocument:
        path = Path(asset.path or "")
        text = _read_text_file(path)
        section = DocumentSection(
            section_id=f"{asset.asset_id}:section:1",
            heading=path.name,
            text=text,
            locator={"path": str(path)},
            section_type="text",
            metadata={"path": str(path)},
        )
        return ParsedDocument(
            document_id=asset.asset_id,
            asset_id=asset.asset_id,
            title=path.name,
            content_type=asset.mime_type or "text/plain",
            plain_text=text,
            sections=[section],
            structural_map={"path": str(path)},
            metadata={"path": str(path), **asset.metadata, "scopes": asset.scope_tags, "source_type": asset.source_type},
        )


class MarkdownAdapter(ParsedDocumentAdapter):
    kind = "markdown"
    supported_suffixes = (".md", ".markdown")

    async def parse(self, asset: KnowledgeAsset) -> ParsedDocument:
        path = Path(asset.path or "")
        text = _read_text_file(path)
        sections: list[DocumentSection] = []
        current_heading = path.stem
        current_lines: list[str] = []
        heading_path: list[str] = []
        section_index = 0
        for line in _tokenize_lines(text):
            heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
            if heading_match:
                if current_lines:
                    section_index += 1
                    sections.append(
                        DocumentSection(
                            section_id=f"{asset.asset_id}:section:{section_index}",
                            heading=current_heading,
                            text="\n".join(current_lines).strip(),
                            locator={"path": str(path), "heading_path": list(heading_path)},
                            section_type="markdown_section",
                            metadata={"path": str(path), "heading_path": list(heading_path)},
                        )
                    )
                    current_lines = []
                current_heading = heading_match.group(2).strip()
                level = len(heading_match.group(1))
                heading_path = heading_path[: max(level - 1, 0)]
                heading_path.append(current_heading)
                continue
            current_lines.append(line)
        if current_lines or not sections:
            section_index += 1
            sections.append(
                DocumentSection(
                    section_id=f"{asset.asset_id}:section:{section_index}",
                    heading=current_heading,
                    text="\n".join(current_lines).strip(),
                    locator={"path": str(path), "heading_path": list(heading_path)},
                    section_type="markdown_section",
                    metadata={"path": str(path), "heading_path": list(heading_path)},
                )
            )
        return ParsedDocument(
            document_id=asset.asset_id,
            asset_id=asset.asset_id,
            title=path.name,
            content_type=asset.mime_type or "text/markdown",
            plain_text=text,
            sections=sections,
            structural_map={"path": str(path), "headings": [section.heading for section in sections if section.heading]},
            metadata={"path": str(path), **asset.metadata, "scopes": asset.scope_tags, "source_type": asset.source_type},
        )


class CodeAdapter(ParsedDocumentAdapter):
    kind = "code"
    supported_suffixes = (".py", ".ts", ".tsx", ".js", ".jsx", ".java", ".go", ".rs", ".cpp", ".c", ".h", ".json", ".yaml", ".yml")

    async def parse(self, asset: KnowledgeAsset) -> ParsedDocument:
        path = Path(asset.path or "")
        text = _read_text_file(path)
        lines = text.splitlines()
        sections: list[DocumentSection] = []
        current_lines: list[str] = []
        current_heading = path.name
        current_symbol = None
        start_line = 1
        symbol_pattern = re.compile(r"^\s*(?:def|class|async def|function|interface|type)\s+([A-Za-z_][A-Za-z0-9_]*)")
        for index, line in enumerate(lines, start=1):
            match = symbol_pattern.match(line)
            if match and current_lines:
                sections.append(
                    DocumentSection(
                        section_id=f"{asset.asset_id}:section:{len(sections) + 1}",
                        heading=current_heading,
                        text="\n".join(current_lines).strip(),
                        locator={"path": str(path), "start_line": start_line, "end_line": index - 1, "symbol": current_symbol},
                        section_type="code_block",
                        metadata={"path": str(path), "symbol": current_symbol},
                    )
                )
                current_lines = []
                start_line = index
            if match:
                current_symbol = match.group(1)
                current_heading = current_symbol
            current_lines.append(line)
        if current_lines:
            sections.append(
                DocumentSection(
                    section_id=f"{asset.asset_id}:section:{len(sections) + 1}",
                    heading=current_heading,
                    text="\n".join(current_lines).strip(),
                    locator={"path": str(path), "start_line": start_line, "end_line": len(lines), "symbol": current_symbol},
                    section_type="code_block",
                    metadata={"path": str(path), "symbol": current_symbol},
                )
            )
        return ParsedDocument(
            document_id=asset.asset_id,
            asset_id=asset.asset_id,
            title=path.name,
            content_type=asset.mime_type or "text/x-code",
            plain_text=text,
            sections=sections,
            structural_map={"path": str(path), "symbols": [section.metadata.get("symbol") for section in sections if section.metadata.get("symbol")]},
            metadata={"path": str(path), **asset.metadata, "scopes": asset.scope_tags, "source_type": asset.source_type},
        )


class PdfAdapter(ParsedDocumentAdapter):
    kind = "pdf"
    supported_suffixes = (".pdf",)

    def _extract_with_pdfplumber(self, path: Path) -> list[str]:
        import pdfplumber  # type: ignore

        pages: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages.append((page.extract_text() or "").strip())
        return pages

    def _extract_with_pypdf(self, path: Path) -> list[str]:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        return [((page.extract_text() or "").strip()) for page in reader.pages]

    def _extract_fallback(self, path: Path) -> list[str]:
        raw = path.read_bytes().decode("latin-1", errors="ignore")
        candidates = re.findall(r"\(([^()]*)\)", raw)
        text = " ".join(item.strip() for item in candidates if item.strip())
        return [text]

    async def parse(self, asset: KnowledgeAsset) -> ParsedDocument:
        path = Path(asset.path or "")
        pages: list[str] = []
        extraction_status = "ok"
        for loader in (self._extract_with_pdfplumber, self._extract_with_pypdf, self._extract_fallback):
            try:
                pages = loader(path)
                if pages:
                    break
            except Exception:
                continue
        if not any(page.strip() for page in pages):
            extraction_status = "needs_ocr"
        sections: list[DocumentSection] = []
        for index, page_text in enumerate(pages or [""], start=1):
            sections.append(
                DocumentSection(
                    section_id=f"{asset.asset_id}:page:{index}",
                    heading=f"Page {index}",
                    text=page_text.strip(),
                    locator={"path": str(path), "page": index, "block_index": 1},
                    section_type="pdf_page",
                    metadata={"path": str(path), "page": index},
                )
            )
        return ParsedDocument(
            document_id=asset.asset_id,
            asset_id=asset.asset_id,
            title=path.name,
            content_type=asset.mime_type or "application/pdf",
            plain_text="\n".join(page.text for page in sections if page.text),
            sections=sections,
            structural_map={"path": str(path), "page_count": len(sections)},
            metadata={
                "path": str(path),
                **asset.metadata,
                "scopes": asset.scope_tags,
                "source_type": asset.source_type,
                "extraction_status": extraction_status,
            },
        )


class DocxAdapter(ParsedDocumentAdapter):
    kind = "docx"
    supported_suffixes = (".docx",)

    def _parse_with_python_docx(self, path: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
        from docx import Document as DocxDocument  # type: ignore

        doc = DocxDocument(str(path))
        paragraphs = []
        heading_path: list[str] = []
        for index, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.startswith("Heading"):
                heading_path = heading_path[:1]
                heading_path.append(text)
            paragraphs.append(
                {"text": text, "heading_path": list(heading_path), "paragraph_index": index, "style": style_name}
            )
        tables = []
        for table_index, table in enumerate(doc.tables):
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        tables.append(
                            {
                                "text": text,
                                "table_index": table_index,
                                "row": row_index,
                                "col": col_index,
                                "heading_path": list(heading_path),
                            }
                        )
        return paragraphs, tables

    def _parse_docx_xml(self, path: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
        paragraphs: list[dict[str, Any]] = []
        tables: list[dict[str, Any]] = []
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml")
        root = ElementTree.fromstring(xml)
        for p_index, para in enumerate(root.findall(".//w:p", ns)):
            texts = [node.text or "" for node in para.findall(".//w:t", ns)]
            text = "".join(texts).strip()
            if text:
                paragraphs.append({"text": text, "heading_path": [], "paragraph_index": p_index, "style": ""})
        for t_index, table in enumerate(root.findall(".//w:tbl", ns)):
            rows = table.findall(".//w:tr", ns)
            for r_index, row in enumerate(rows):
                cells = row.findall(".//w:tc", ns)
                for c_index, cell in enumerate(cells):
                    texts = [node.text or "" for node in cell.findall(".//w:t", ns)]
                    text = "".join(texts).strip()
                    if text:
                        tables.append({"text": text, "table_index": t_index, "row": r_index, "col": c_index, "heading_path": []})
        return paragraphs, tables

    async def parse(self, asset: KnowledgeAsset) -> ParsedDocument:
        path = Path(asset.path or "")
        paragraphs: list[dict[str, Any]] = []
        tables: list[dict[str, Any]] = []
        for loader in (self._parse_with_python_docx, self._parse_docx_xml):
            try:
                paragraphs, tables = loader(path)
                if paragraphs or tables:
                    break
            except Exception:
                continue
        sections: list[DocumentSection] = []
        for para in paragraphs:
            sections.append(
                DocumentSection(
                    section_id=f"{asset.asset_id}:paragraph:{len(sections) + 1}",
                    heading=(para.get("heading_path") or [path.stem])[-1] if para.get("heading_path") else path.stem,
                    text=para["text"],
                    locator={
                        "path": str(path),
                        "heading_path": list(para.get("heading_path") or []),
                        "paragraph_index": para.get("paragraph_index"),
                    },
                    section_type="docx_paragraph",
                    metadata={"path": str(path), "heading_path": list(para.get("heading_path") or [])},
                )
            )
        for table in tables:
            sections.append(
                DocumentSection(
                    section_id=f"{asset.asset_id}:table:{len(sections) + 1}",
                    heading="Table Cell",
                    text=table["text"],
                    locator={
                        "path": str(path),
                        "heading_path": list(table.get("heading_path") or []),
                        "table_index": table.get("table_index"),
                        "row": table.get("row"),
                        "col": table.get("col"),
                    },
                    section_type="docx_table_cell",
                    metadata={"path": str(path), "heading_path": list(table.get("heading_path") or [])},
                )
            )
        plain_text = "\n".join(section.text for section in sections if section.text)
        return ParsedDocument(
            document_id=asset.asset_id,
            asset_id=asset.asset_id,
            title=path.name,
            content_type=asset.mime_type or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            plain_text=plain_text,
            sections=sections,
            structural_map={"path": str(path), "paragraph_count": len(paragraphs), "table_cell_count": len(tables)},
            metadata={"path": str(path), **asset.metadata, "scopes": asset.scope_tags, "source_type": asset.source_type},
        )


class MemoryRecordAdapter(ParsedDocumentAdapter):
    kind = "memory_record"
    supported_suffixes = ()

    async def parse(self, asset: KnowledgeAsset) -> ParsedDocument:
        content = str(asset.metadata.get("content", ""))
        locator = {key: asset.metadata.get(key) for key in ("record_id", "thread_id", "task_id") if asset.metadata.get(key) is not None}
        section = DocumentSection(
            section_id=f"{asset.asset_id}:section:1",
            heading=asset.metadata.get("kind", asset.asset_id),
            text=content,
            locator=locator,
            section_type="memory_record",
            metadata=dict(asset.metadata),
        )
        return ParsedDocument(
            document_id=asset.asset_id,
            asset_id=asset.asset_id,
            title=asset.metadata.get("kind", asset.asset_id),
            content_type=asset.mime_type or "application/x-memory-record",
            plain_text=content,
            sections=[section],
            structural_map=locator,
            metadata={**asset.metadata, "scopes": asset.scope_tags, "source_type": asset.source_type},
        )


class WorkspaceArtifactAdapter(ParsedDocumentAdapter):
    kind = "workspace_artifact"
    supported_suffixes = ()

    async def parse(self, asset: KnowledgeAsset) -> ParsedDocument:
        content = asset.metadata.get("content", "")
        if not isinstance(content, str):
            content = str(content)
        locator = {key: asset.metadata.get(key) for key in ("artifact_id", "task_id", "owner_agent") if asset.metadata.get(key) is not None}
        section = DocumentSection(
            section_id=f"{asset.asset_id}:section:1",
            heading=asset.metadata.get("name", asset.asset_id),
            text=content,
            locator=locator,
            section_type="workspace_artifact",
            metadata=dict(asset.metadata),
        )
        return ParsedDocument(
            document_id=asset.asset_id,
            asset_id=asset.asset_id,
            title=asset.metadata.get("name", asset.asset_id),
            content_type=asset.mime_type or "application/x-workspace-artifact",
            plain_text=content,
            sections=[section],
            structural_map=locator,
            metadata={**asset.metadata, "scopes": asset.scope_tags, "source_type": asset.source_type},
        )


DEFAULT_ADAPTERS: tuple[type[ParsedDocumentAdapter], ...] = (
    MarkdownAdapter,
    TextAdapter,
    CodeAdapter,
    PdfAdapter,
    DocxAdapter,
    MemoryRecordAdapter,
    WorkspaceArtifactAdapter,
)


def infer_asset_from_path(path: str | Path, *, scopes: Iterable[str] | None = None, source_type: str = "filesystem") -> KnowledgeAsset:
    path_obj = Path(path)
    mime_type, _ = mimetypes.guess_type(path_obj.name)
    return KnowledgeAsset(
        asset_id=path_obj.as_posix(),
        source_type=source_type,
        path=str(path_obj),
        uri=str(path_obj),
        mime_type=mime_type or "application/octet-stream",
        scope_tags=list(scopes or []),
        metadata={"path": str(path_obj), "suffix": path_obj.suffix.lower()},
    )
